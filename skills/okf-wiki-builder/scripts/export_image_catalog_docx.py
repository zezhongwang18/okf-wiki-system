#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import os
import re
import tempfile
from pathlib import Path
from typing import Optional


IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff", ".webp"}


def require_docx():
    try:
        from docx import Document  # type: ignore
        from docx.shared import Inches, Pt  # type: ignore
    except Exception as exc:
        raise SystemExit(
            "This exporter requires python-docx. Run it with the Codex bundled Python runtime "
            "or install python-docx in the active Python environment."
        ) from exc
    return Document, Inches, Pt


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def frontmatter_value(markdown: str, key: str) -> str:
    if not markdown.startswith("---\n"):
        return ""
    end = markdown.find("\n---\n", 4)
    if end == -1:
        return ""
    yaml = markdown[4:end]
    match = re.search(rf"^{re.escape(key)}:\s*(.+)$", yaml, re.MULTILINE)
    return match.group(1).strip().strip("\"'") if match else ""


def frontmatter_list(markdown: str, key: str) -> list[str]:
    if not markdown.startswith("---\n"):
        return []
    end = markdown.find("\n---\n", 4)
    if end == -1:
        return []
    lines = markdown[4:end].splitlines()
    values: list[str] = []
    in_key = False
    for line in lines:
        if re.match(rf"^{re.escape(key)}:\s*$", line):
            in_key = True
            continue
        if in_key:
            item = re.match(r"^\s*-\s*(.+)$", line)
            if item:
                values.append(item.group(1).strip().strip("\"'"))
                continue
            if re.match(r"^[A-Za-z0-9_-]+:", line):
                break
    return values


def section(markdown: str, heading: str) -> str:
    pattern = re.compile(rf"^# {re.escape(heading)}\s*$([\s\S]*?)(?=^# |\Z)", re.MULTILINE)
    match = pattern.search(markdown)
    return match.group(1).strip() if match else ""


def subsection(parent: str, heading: str) -> str:
    pattern = re.compile(rf"^## {re.escape(heading)}\s*$([\s\S]*?)(?=^## |\Z)", re.MULTILINE)
    match = pattern.search(parent)
    return match.group(1).strip() if match else ""


def clean_lines(markdown_block: str) -> list[str]:
    lines: list[str] = []
    for raw in markdown_block.splitlines():
        line = raw.strip()
        if not line:
            continue
        line = re.sub(r"^[-*]\s+", "", line)
        line = re.sub(r"`([^`]+)`", r"\1", line)
        line = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line)
        lines.append(line)
    return lines


def first_nonempty(*values: str) -> str:
    for value in values:
        if value and value.strip() and not value.strip().lower().startswith("todo"):
            return value.strip()
    return ""


def trim_text(text: str, limit: int = 600) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) <= limit:
        return text
    return text[:limit].rstrip() + " ..."


def normalize_caption(caption: str) -> str:
    caption = caption.strip()
    caption = re.sub(r"^(图|圖|figure|fig\.)\s*[\d一二三四五六七八九十]+[：:.\-\s]*", "", caption, flags=re.I)
    return caption.strip() or "未命名图片"


def extract_context_summary(source_context: str) -> str:
    same_heading = subsection(source_context, "Same-Heading Text Before Image")
    recent = subsection(source_context, "Recent Text Before Image")
    lines = clean_lines(same_heading or recent)
    return trim_text(" ".join(lines), 700) if lines else "未提取"


def extract_field(source_context: str, field: str) -> str:
    match = re.search(rf"^- {re.escape(field)}:\s*(.+)$", source_context, re.MULTILINE)
    return match.group(1).strip() if match else "未提取"


def convert_for_word(image_path: Path, temp_dir: Path) -> Optional[Path]:
    if image_path.suffix.lower() in {".png", ".jpg", ".jpeg", ".gif"}:
        return image_path
    if image_path.suffix.lower() not in IMAGE_SUFFIXES:
        return None
    try:
        from PIL import Image  # type: ignore

        converted = temp_dir / f"{image_path.stem}.png"
        with Image.open(image_path) as image:
            if image.mode not in {"RGB", "RGBA"}:
                image = image.convert("RGBA")
            image.save(converted, "PNG")
        return converted
    except Exception:
        return None


def add_label(doc, label: str, value: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(label)
    run.bold = True
    paragraph.add_run(value or "未提取")


def set_default_font(document, Pt) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)


def build_catalog(bundle: Path, output: Path, max_width_inches: float) -> int:
    Document, Inches, Pt = require_docx()
    assets_dir = bundle / "assets"
    raw_root = bundle
    pages = sorted(path for path in assets_dir.glob("asset-*.md") if path.is_file())
    if not pages:
        raise SystemExit(f"No asset pages found under {assets_dir}")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_default_font(doc, Pt)
    doc.add_heading("图片素材汇总", level=0)
    doc.add_paragraph("本文件汇总 OKF Wiki 摄入过程中整理出的图片资产。每张图片包含图题、来源、上下文摘要、OCR/可见文字和图片本体。")

    added = 0
    manifest: list[dict] = []
    with tempfile.TemporaryDirectory() as tmp:
        temp_dir = Path(tmp)
        for page in pages:
            markdown = read_text(page)
            resource = frontmatter_value(markdown, "resource")
            if not resource:
                continue
            image_path = raw_root / resource
            if not image_path.exists():
                continue
            word_image = convert_for_word(image_path, temp_dir)

            title = frontmatter_value(markdown, "title") or image_path.stem
            description = section(markdown, "Description")
            applicable_questions = section(markdown, "Applicable Questions")
            not_applicable_to = section(markdown, "Not Applicable To")
            source_context = section(markdown, "Source Context")
            visible_text = section(markdown, "Visible Text")
            visual_notes = section(markdown, "Visual Notes")
            sources = frontmatter_list(markdown, "sources")

            caption = first_nonempty(
                extract_field(source_context, "Caption"),
                title,
            )
            caption = normalize_caption(caption)
            summary = first_nonempty(
                description,
                visual_notes,
                extract_context_summary(source_context),
            )
            added += 1
            doc.add_page_break() if added > 1 else None
            doc.add_heading(f"图 {added}：{caption}", level=1)

            add_label(doc, "检索摘要：", trim_text(summary, 500))
            add_label(doc, "来源文件：", extract_field(source_context, "Source file"))
            add_label(doc, "来源页面：", ", ".join(sources) if sources else "未提取")
            add_label(doc, "适用问题：", trim_text(" ".join(clean_lines(applicable_questions)), 700) if applicable_questions else "Not assigned.")
            add_label(doc, "不适用问题：", trim_text(" ".join(clean_lines(not_applicable_to)), 700) if not_applicable_to else "Not assigned.")
            add_label(doc, "绑定标题：", extract_field(source_context, "Nearest heading"))
            add_label(doc, "绑定置信度：", extract_field(source_context, "Binding confidence"))
            add_label(doc, "上下文摘要：", extract_context_summary(source_context))
            add_label(doc, "OCR/可见文字：", trim_text(" ".join(clean_lines(visible_text)), 700) if visible_text else "未提取")

            doc.add_paragraph("图片本体：").runs[0].bold = True
            if word_image:
                doc.add_picture(str(word_image), width=Inches(max_width_inches))
                manifest.append({
                    "index": added,
                    "asset_page": page.relative_to(bundle).as_posix(),
                    "resource": resource,
                    "caption": caption,
                    "source_file": extract_field(source_context, "Source file"),
                    "applicable_questions": clean_lines(applicable_questions),
                    "not_applicable_to": clean_lines(not_applicable_to),
                })
            else:
                doc.add_paragraph(f"无法嵌入该图片格式：{image_path.name}")

    if added == 0:
        raise SystemExit("No embeddable asset images were found.")
    doc.save(output)
    output.with_suffix(".manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    return added


def main() -> None:
    if os.environ.get("OKF_FINALIZE_RUNNING") != "1":
        raise SystemExit(
            "Do not run export_image_catalog_docx.py directly as a completion step. "
            "Run finalize_bundle.py so image catalog export, upload export, and validation happen together."
        )
    parser = argparse.ArgumentParser(description="Export OKF image assets into a Word catalog with embedded image bodies.")
    parser.add_argument("bundle", nargs="?", default=".", help="OKF bundle root")
    parser.add_argument("--output", help="Output .docx path; defaults to bundle/exports/image-catalog.docx")
    parser.add_argument("--max-width-inches", type=float, default=5.8, help="Maximum rendered image width in the Word document")
    args = parser.parse_args()

    bundle = Path(args.bundle).resolve()
    output = Path(args.output).resolve() if args.output else bundle / "exports" / "image-catalog.docx"
    count = build_catalog(bundle, output, args.max_width_inches)
    print(f"Wrote {output} with {count} embedded image item(s).")


if __name__ == "__main__":
    main()
